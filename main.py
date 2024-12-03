import os
import openpyxl
from docx import Document


def extract_excel_data_to_word(folder_path):
    # Ensure the folder path is valid
    if not os.path.exists(folder_path):
        print("Invalid folder path.")
        return

    # Create a new folder for the summary file
    summary_folder = os.path.join(folder_path, "Summary_Output")
    os.makedirs(summary_folder, exist_ok=True)

    # Initialize Word document
    doc = Document()
    doc.add_heading("Summary of Extracted Data", level=1)

    # Create a table in the Word document
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Add table headers
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Excel File Name'
    header_cells[1].text = 'Extracted Value from D2'

    # Get all Excel files sorted in descending order
    excel_files = sorted([f for f in os.listdir(folder_path) if f.endswith(".xlsx") or f.endswith(".xlsm")],
                         reverse=True)

    # Track if any valid data is written
    data_written = False

    # Process each Excel file
    for file_name in excel_files:
        file_path = os.path.join(folder_path, file_name)
        try:
            # Open the Excel file
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            if "Validation" in workbook.sheetnames:
                sheet = workbook["Validation"]
                # Extract value from cell D2
                cell_value = sheet["D2"].value
                if cell_value is not None:
                    # Add row to the table for each valid file
                    row_cells = table.add_row().cells
                    row_cells[0].text = file_name
                    row_cells[1].text = str(cell_value)
                    data_written = True
                else:
                    print(f"{file_name}: Cell D2 is empty. Skipping.")
            else:
                print(f"{file_name}: Sheet 'Validation' not found. Skipping.")
        except Exception as e:
            print(f"{file_name}: Error reading file ({e}). Skipping.")

    # Save the Word document only if data was written
    if data_written:
        summary_path = os.path.join(summary_folder, "Summary.docx")
        doc.save(summary_path)
        print(f"Summary saved at {summary_path}")
    else:
        print("No valid data found. Summary document was not created.")


# Use the specified folder path
folder_path = r"D:\1st batch Investment Reports WithID to Attain\Rejected Returns Reasons"
extract_excel_data_to_word(folder_path)
